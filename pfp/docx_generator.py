"""
DOCX 생성 모듈
LENS 템플릿 스타일로 Claude가 쓴 내용을 새로 채우는 방식
블랙스카이 내용 복사 없음
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── LENS 색상 팔레트 ──────────────────────────────────────
DARK_NAVY  = "1F3864"
BLUE       = "2E75B6"
GREEN      = "375623"
LIGHT_BLUE = "DEEAF1"
LIGHT_GRAY = "F5F5F5"
WHITE      = "FFFFFF"
GRAY_TEXT  = "404040"


def generate_docx(sections: dict, ticker: str, company_name: str, output_dir: str) -> str:
    raw = sections.get("_raw", "")
    parsed = _parse_sections(raw)

    doc = Document()
    _set_page(doc)
    _set_default_font(doc)

    # 커버 페이지
    _add_cover(doc, ticker, company_name, raw)
    doc.add_page_break()

    # 섹션별 내용
    _add_sections(doc, raw)

    filename = f"{ticker}_LENS_Report.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"[DOCX] 저장 완료: {filepath}")
    return filepath


def _set_page(doc):
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin  = sec.bottom_margin = Inches(0.8)


def _set_default_font(doc):
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)


def _rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _add_bottom_border(paragraph, color=BLUE, size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_header_info(raw: str) -> dict:
    """raw 텍스트에서 헤더 정보 추출"""
    info = {
        'opinion': 'BUY', 'price': '$0.00', 'mktcap': '$0B',
        'bull_target': '$0', 'bear_target': '$0',
        'highlights': [], 'exchange': 'NYSE', 'sector': '-',
        'ceo': '-', 'hq': '-', 'slogan': ''
    }
    lines = raw.split('\n')
    for i, line in enumerate(lines):
        l = line.strip().lower()
        if '투자의견' in line and ('buy' in line.upper() or 'hold' in line.upper() or 'sell' in line.upper()):
            for op in ['BUY','HOLD','SELL']:
                if op in line.upper():
                    info['opinion'] = op
                    break
        if '현재주가' in line and '$' in line:
            m = re.search(r'\$[\d,]+\.?\d*', line)
            if m: info['price'] = m.group()
        if '시가총액' in line and '$' in line:
            m = re.search(r'\$[\d,.]+[BMT]?', line)
            if m: info['mktcap'] = m.group()
        if 'bull' in l and '목표주가' in line:
            m = re.search(r'\$[\d,]+', line)
            if m: info['bull_target'] = m.group()
        if 'bear' in l and '목표주가' in line:
            m = re.search(r'\$[\d,]+', line)
            if m: info['bear_target'] = m.group()
        if line.strip().startswith('•') or line.strip().startswith('*•'):
            hl = line.strip().lstrip('*•').strip()
            if hl and len(info['highlights']) < 5:
                info['highlights'].append(hl)
        if 'ceo' in l and '|' in line:
            parts = line.split('|')
            for j, p in enumerate(parts):
                if 'CEO' in p and j+1 < len(parts):
                    info['ceo'] = parts[j+1].strip()
        if '거래소' in line and '|' in line:
            parts = line.split('|')
            for j, p in enumerate(parts):
                if '거래소' in p and j+1 < len(parts):
                    info['exchange'] = parts[j+1].strip()
        if '업종' in line and '|' in line:
            parts = line.split('|')
            for j, p in enumerate(parts):
                if '업종' in p and j+1 < len(parts):
                    info['sector'] = parts[j+1].strip()
        if '본사' in line and '|' in line:
            parts = line.split('|')
            for j, p in enumerate(parts):
                if '본사' in p and j+1 < len(parts):
                    info['hq'] = parts[j+1].strip()
    return info


def _add_cover(doc, ticker, company_name, raw):
    """LENS 스타일 커버 페이지"""
    info = _parse_header_info(raw)

    # ── 상단 헤더 줄 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run("[ L ]  LENS CAPITAL RESEARCH")
    r1.font.name = 'Calibri'; r1.font.bold = True
    r1.font.size = Pt(16); r1.font.color.rgb = _rgb(DARK_NAVY)
    r2 = p.add_run("   EQUITY RESEARCH")
    r2.font.name = 'Calibri'; r2.font.size = Pt(12)
    r2.font.color.rgb = _rgb(BLUE)

    # ── 블루 구분선 ──
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(6)
    _add_bottom_border(p2, BLUE, "12")

    # ── EQUITY RESEARCH REPORT ──
    p3 = doc.add_paragraph()
    r = p3.add_run("EQUITY RESEARCH REPORT")
    r.font.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = _rgb(BLUE)
    p3.paragraph_format.space_after = Pt(4)

    # ── 티커 + 회사명 ──
    p4 = doc.add_paragraph()
    r1 = p4.add_run(ticker)
    r1.font.bold = True; r1.font.size = Pt(24)
    r1.font.color.rgb = _rgb(DARK_NAVY)
    r2 = p4.add_run(f"   {company_name}")
    r2.font.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = _rgb(GRAY_TEXT)
    p4.paragraph_format.space_after = Pt(4)

    # ── 슬로건 ──
    slogan = _extract_slogan(raw, company_name)
    p5 = doc.add_paragraph()
    r = p5.add_run(slogan)
    r.font.italic = True; r.font.size = Pt(12)
    r.font.color.rgb = _rgb(DARK_NAVY)
    p5.paragraph_format.space_after = Pt(8)

    # ── BUY/PRICE/HIGHLIGHTS 3열 표 ──
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 열 너비
    widths = [Inches(1.4), Inches(2.8), Inches(3.0)]
    for i, cell in enumerate(tbl.rows[0].cells):
        cell.width = widths[i]

    # 셀1: BUY/HOLD/SELL
    c1 = tbl.rows[0].cells[0]
    _set_cell_bg(c1, GREEN)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(info['opinion'])
    r.font.bold = True; r.font.size = Pt(28)
    r.font.color.rgb = _rgb(WHITE)

    # 셀2: 주가/시총
    c2 = tbl.rows[0].cells[1]
    _set_cell_bg(c2, LIGHT_BLUE)
    p = c2.paragraphs[0]
    r1 = p.add_run(f"{info['price']}")
    r1.font.bold = True; r1.font.size = Pt(18)
    r1.font.color.rgb = _rgb(DARK_NAVY)
    p2 = c2.add_paragraph(f"{info['mktcap']}")
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].font.color.rgb = _rgb(GRAY_TEXT)
    p3 = c2.add_paragraph(f"Bull: {info['bull_target']}  /  Bear: {info['bear_target']}")
    p3.runs[0].font.size = Pt(9)
    p3.runs[0].font.color.rgb = _rgb(BLUE)

    # 셀3: KEY HIGHLIGHTS
    c3 = tbl.rows[0].cells[2]
    ph = c3.paragraphs[0]
    rh = ph.add_run("KEY HIGHLIGHTS")
    rh.font.bold = True; rh.font.size = Pt(9)
    rh.font.color.rgb = _rgb(DARK_NAVY)
    for hl in info['highlights'][:5]:
        pb = c3.add_paragraph()
        rb = pb.add_run(f"• {hl[:80]}")
        rb.font.size = Pt(8)
        rb.font.color.rgb = _rgb(GRAY_TEXT)
        pb.paragraph_format.space_before = Pt(1)
        pb.paragraph_format.space_after  = Pt(1)

    doc.add_paragraph()

    # ── 기업 정보 6열 표 ──
    tbl2 = doc.add_table(rows=2, cols=6)
    tbl2.style = 'Table Grid'
    headers = ['거래소','업종','CEO','발행일','본사','애널리스트']
    values  = [
        info['exchange'], info['sector'][:20], info['ceo'],
        'Apr. 2026', info['hq'][:20], 'Lens AI'
    ]
    col_w = Inches(1.2)
    for i, (h, v) in enumerate(zip(headers, values)):
        hc = tbl2.rows[0].cells[i]
        vc = tbl2.rows[1].cells[i]
        hc.width = vc.width = col_w
        _set_cell_bg(hc, DARK_NAVY)
        ph = hc.paragraphs[0]
        rh = ph.add_run(h)
        rh.font.bold = True; rh.font.size = Pt(8)
        rh.font.color.rgb = _rgb(WHITE)
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pv = vc.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(8)
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 푸터 슬로건 ──
    pf = doc.add_paragraph()
    rf = pf.add_run('"선명하게, 다르게 본다."  Insight in focus.   [ L ]  LENS CAPITAL RESEARCH  │  lens-capital.io')
    rf.font.italic = True; rf.font.size = Pt(8)
    rf.font.color.rgb = _rgb(DARK_NAVY)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(4)

    doc.add_paragraph()

    # ── 목차 ──
    _add_section_heading(doc, "01  목차 (Table of Contents)")
    toc = [
        "II.   투자의견 요약 (Executive Summary)",
        "III.  기업 개요 (Company Overview)",
        "IV.  시황·업종·산업 분석 (Top-down Analysis)",
        "V.   투자 근거 (Investment Points)",
        "VI.  재무제표 & KPI 분석 (Financials)",
        "VII. 동종업체 비교 분석 (Peer Analysis)",
        "VIII. 리스크 요인 (Risk Factors)",
        "IX.  적정주가 산출 (Valuation)",
        "X.   종합 결론 (Conclusion)",
    ]
    for t in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(t)
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb(GRAY_TEXT)


def _extract_slogan(raw, company_name):
    for line in raw.split('\n'):
        line = line.strip().strip('*').strip()
        if line.startswith('[') or not line:
            continue
        if ('—' in line or '-' in line) and len(line) > 20 and len(line) < 120:
            if any(w in line for w in ['성장', '전환', '수요', '선두', '변곡', '기회', '혁신', 'AI', '글로벌']):
                return line
    return f"{company_name} — 핵심 투자 포인트"


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    _add_bottom_border(p, BLUE, "8")
    r = p.add_run(text)
    r.font.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = _rgb(DARK_NAVY)


def _add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = _rgb(BLUE)


def _add_sections(doc, raw):
    lines = raw.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('## '):
            heading = stripped[3:].strip()
            # HEADER 섹션은 커버에서 이미 처리
            if 'header' not in heading.upper() and 'HEADER' not in heading:
                _add_section_heading(doc, heading)
            i += 1
            continue

        # 마크다운 표
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # 소제목 ###
        if stripped.startswith('### '):
            _add_sub_heading(doc, stripped[4:])
            i += 1
            continue

        # **굵은 제목** (단독 줄)
        if re.match(r'^\*\*[^*]+\*\*$', stripped) and len(stripped) > 4:
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip('*'))
            r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = _rgb(DARK_NAVY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # 📌 강조
        if stripped.startswith('📌'):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = _rgb(BLUE)
            p.paragraph_format.space_before = Pt(6)
            i += 1
            continue

        # 🔴🟡🟢 리스크
        if stripped.startswith(('🔴','🟡','🟢')):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.bold = True; r.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
            i += 1
            continue

        # ▸ 불릿
        if stripped.startswith(('▸','*▸')):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            clean = stripped.lstrip('*▸ ')
            _add_mixed_run(p, clean, Pt(10))
            i += 1
            continue

        # • 불릿
        if stripped.startswith('• ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _add_mixed_run(p, stripped[2:], Pt(10))
            i += 1
            continue

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 일반 본문
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _add_mixed_run(p, stripped, Pt(10))
        i += 1


def _add_mixed_run(paragraph, text, size):
    """**볼드** 텍스트 파싱"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.font.bold = True
        else:
            r = paragraph.add_run(part)
        r.font.size = size
        r.font.color.rgb = _rgb(GRAY_TEXT)


def _add_table(doc, lines):
    """마크다운 표 → DOCX 표"""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue
        cells = [c.strip().strip('*') for c in line.strip().strip('|').split('|')]
        if cells:
            rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    rows = [r + [''] * (col_count - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = 'Table Grid'

    page_width = 9360  # DXA (US Letter - 1" margins each side)
    col_w = page_width // col_count

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Pt(col_w)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(val)
            if i == 0:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(WHITE)
                _set_cell_bg(cell, DARK_NAVY)
            else:
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(GRAY_TEXT)
                if i % 2 == 0:
                    _set_cell_bg(cell, LIGHT_GRAY)

    doc.add_paragraph()


def _parse_sections(raw: str) -> dict:
    sections = {}
    current = 'header'
    buf = []
    for line in raw.split('\n'):
        if line.startswith('## '):
            if buf: sections[current] = '\n'.join(buf).strip()
            current = line[3:].strip().lower().replace(' ', '_')
            buf = []
        else:
            buf.append(line)
    if buf: sections[current] = '\n'.join(buf).strip()
    sections['_raw'] = raw
    return sections


# ── PDF 변환 ──────────────────────────────────────────────────────────────────

def generate_pdf(sections: dict, ticker: str, company_name: str, output_dir: str) -> str:
    """DOCX 생성 후 PDF로 변환해서 반환"""
    import subprocess, sys, platform

    # 1. 먼저 DOCX 생성
    docx_path = generate_docx(sections, ticker, company_name, output_dir)
    pdf_path  = docx_path.replace(".docx", ".pdf")

    system = platform.system()

    # 2. LibreOffice (Windows/Mac/Linux 공통)
    lo_cmds = []
    if system == "Windows":
        lo_cmds = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
    elif system == "Darwin":
        lo_cmds = ["/Applications/LibreOffice.app/Contents/MacOS/soffice"]
    else:
        lo_cmds = ["libreoffice", "soffice"]

    for lo in lo_cmds:
        try:
            result = subprocess.run(
                [lo, "--headless", "--convert-to", "pdf",
                 "--outdir", output_dir, docx_path],
                capture_output=True, timeout=60
            )
            if result.returncode == 0 and os.path.exists(pdf_path):
                print(f"[PDF] LibreOffice 변환 완료: {pdf_path}")
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # 3. docx2pdf 라이브러리 시도
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            print(f"[PDF] docx2pdf 변환 완료: {pdf_path}")
            return pdf_path
    except ImportError:
        pass
    except Exception as e:
        print(f"[PDF] docx2pdf 오류: {e}")

    # 4. 변환 실패 시 DOCX 경로 반환 (폴백)
    print("[PDF] 변환 실패 — DOCX로 대체. LibreOffice 또는 docx2pdf 설치 필요.")
    return docx_path

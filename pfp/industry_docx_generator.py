"""
industry_docx_generator.py — 산업 리서치 레포트 DOCX/PDF 생성기
LENS Industry Research 스타일 (Space/AI Connectivity 템플릿 기준)
"""
import os, re, subprocess, platform
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── LENS Industry 색상 팔레트 ──────────────────────────────────────────────────
BRAND_RED    = "C1440D"
DARK         = "1A1A1A"
DARK2        = "222222"
GRAY         = "5F6368"
GRAY2        = "333333"
LIGHT_GRAY   = "F5F5F5"
WHITE        = "FFFFFF"
OW_GREEN     = "2E7D32"   # OVERWEIGHT
NT_BLUE      = "1565C0"   # NEUTRAL
UW_RED       = "B71C1C"   # UNDERWEIGHT
DARK_NAVY    = "1F3864"


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_bottom_border(paragraph, color: str, size: str = "16", space: str = "6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 헤더 파싱 ─────────────────────────────────────────────────────────────────

def _parse_industry_header(raw: str) -> dict:
    info = {
        "opinion": "OVERWEIGHT",
        "return_12m": "+XX% [E]",
        "bull": "+XX%",
        "bear": "-XX%",
        "tagline": "",
        "benchmark": "",
        "coverage": "",
        "highlights": [],
    }
    header_section = ""
    in_header = False
    for line in raw.split("\n"):
        if line.strip() == "## HEADER":
            in_header = True
            continue
        if in_header and line.startswith("## "):
            break
        if in_header:
            header_section += line + "\n"

    for line in header_section.split("\n"):
        l = line.strip()
        if l.lower().startswith("의견:"):
            info["opinion"] = l.split(":", 1)[1].strip()
        elif l.lower().startswith("12m수익률:"):
            info["return_12m"] = l.split(":", 1)[1].strip()
        elif l.lower().startswith("bull:"):
            info["bull"] = l.split(":", 1)[1].strip()
        elif l.lower().startswith("bear:"):
            info["bear"] = l.split(":", 1)[1].strip()
        elif l.startswith("슬로건:"):
            info["tagline"] = l.split(":", 1)[1].strip()
        elif l.startswith("벤치마크:"):
            info["benchmark"] = l.split(":", 1)[1].strip()
        elif l.startswith("커버리지:"):
            info["coverage"] = l.split(":", 1)[1].strip()
        elif re.match(r"KEY_HIGHLIGHT_\d+:", l):
            hl = l.split(":", 1)[1].strip()
            if hl:
                info["highlights"].append(hl)
    return info


# ── 커버 페이지 ───────────────────────────────────────────────────────────────

def _add_industry_cover(doc, name_kr: str, name_en: str, meta_dict: dict, info: dict):
    today = datetime.now().strftime("%Y.%m.%d")

    # 1) 브랜드 헤더 줄
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("[ L ] ")
    r1.font.name = "Arial"; r1.font.bold = True
    r1.font.size = Pt(15); r1.font.color.rgb = _rgb(BRAND_RED)
    r2 = p.add_run("LENS CAPITAL RESEARCH   ")
    r2.font.name = "Arial"; r2.font.bold = True
    r2.font.size = Pt(15); r2.font.color.rgb = _rgb(DARK2)
    r3 = p.add_run("INDUSTRY RESEARCH")
    r3.font.name = "Arial"; r3.font.size = Pt(10)
    r3.font.color.rgb = _rgb(GRAY)

    # 2) 레드 구분선
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(10)
    _add_bottom_border(p2, BRAND_RED, "16", "6")
    p2.add_run(" ").font.size = Pt(1)

    # 3) "INDUSTRY RESEARCH REPORT" 레이블
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    r = p3.add_run("INDUSTRY RESEARCH REPORT")
    r.font.name = "Arial"; r.font.bold = True
    r.font.size = Pt(9); r.font.color.rgb = _rgb(BRAND_RED)

    # 4) 산업명 (한국어 크게 + 영어 작게)
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(4)
    r1 = p4.add_run(name_kr + "  ")
    r1.font.name = "Arial"; r1.font.bold = True
    r1.font.size = Pt(20); r1.font.color.rgb = _rgb(DARK)
    r2 = p4.add_run(name_en)
    r2.font.name = "Arial"; r2.font.size = Pt(11)
    r2.font.color.rgb = _rgb(GRAY)

    # 5) 태그라인
    tagline = info.get("tagline") or meta_dict.get("tagline", "")
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(12)
    r = p5.add_run(tagline + " | " + today)
    r.font.name = "Arial"; r.font.italic = True
    r.font.size = Pt(10.5); r.font.color.rgb = _rgb(GRAY2)

    # 6) 3열 요약 표: 의견 | 수익률/Bull·Bear | KEY HIGHLIGHTS
    opinion = info.get("opinion", "OVERWEIGHT").upper()
    badge_color = OW_GREEN if "OVER" in opinion else (NT_BLUE if "NEUTRAL" in opinion else UW_RED)
    badge_label = {"OVERWEIGHT": "OVERWEIGHT", "NEUTRAL": "NEUTRAL", "UNDERWEIGHT": "UNDERWEIGHT"}.get(
        opinion.split()[0], opinion
    )

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    widths = [Inches(1.4), Inches(2.2), Inches(3.6)]
    for i, cell in enumerate(tbl.rows[0].cells):
        cell.width = widths[i]

    # 셀1: Opinion badge
    c1 = tbl.rows[0].cells[0]
    _set_cell_bg(c1, badge_color)
    pb = c1.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pb.add_run(badge_label)
    r.font.name = "Arial"; r.font.bold = True
    r.font.size = Pt(13); r.font.color.rgb = _rgb(WHITE)
    p_sub = c1.add_paragraph("산업 의견")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.name = "Arial"
    p_sub.runs[0].font.size = Pt(8)
    p_sub.runs[0].font.color.rgb = _rgb("E8F5E9" if "OVER" in opinion else "E3F2FD" if "NEUTRAL" in opinion else "FFEBEE")

    # 셀2: 수익률/Bull·Bear
    c2 = tbl.rows[0].cells[1]
    _set_cell_bg(c2, "F8F9FA")
    p_ret_label = c2.paragraphs[0]
    r_label = p_ret_label.add_run("12M 산업지수 수익률")
    r_label.font.name = "Arial"; r_label.font.size = Pt(8)
    r_label.font.color.rgb = _rgb(GRAY)
    p_ret_val = c2.add_paragraph(info.get("return_12m", "+XX% [E]"))
    p_ret_val.runs[0].font.name = "Arial"; p_ret_val.runs[0].font.bold = True
    p_ret_val.runs[0].font.size = Pt(16)
    p_ret_val.runs[0].font.color.rgb = _rgb(DARK)
    p_bull_bear = c2.add_paragraph(f"Bull {info.get('bull', '+XX%')} / Bear {info.get('bear', '-XX%')}")
    p_bull_bear.runs[0].font.name = "Arial"; p_bull_bear.runs[0].font.size = Pt(9)
    p_bull_bear.runs[0].font.color.rgb = _rgb(GRAY)
    p_hl_label = c2.add_paragraph("KEY HIGHLIGHTS")
    p_hl_label.runs[0].font.name = "Arial"; p_hl_label.runs[0].font.bold = True
    p_hl_label.runs[0].font.size = Pt(8)
    p_hl_label.runs[0].font.color.rgb = _rgb(DARK_NAVY)

    # 셀3: Key Highlights 목록
    c3 = tbl.rows[0].cells[2]
    for j, hl in enumerate(info.get("highlights", [])[:5]):
        txt = f"• {hl[:90]}"
        if j == 0:
            p_hl = c3.paragraphs[0]
        else:
            p_hl = c3.add_paragraph()
        p_hl.paragraph_format.space_before = Pt(2)
        p_hl.paragraph_format.space_after = Pt(2)
        r_hl = p_hl.add_run(txt)
        r_hl.font.name = "Arial"; r_hl.font.size = Pt(8.5)
        r_hl.font.color.rgb = _rgb(GRAY2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 7) 메타 정보 행: 벤치마크 | 커버리지 | 발행일 | 애널리스트
    coverage_tickers = info.get("coverage") or meta_dict.get("coverage", "")
    n_tickers = len([t.strip() for t in coverage_tickers.split(",") if t.strip()])
    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.style = "Table Grid"
    headers = ["벤치마크", "커버리지 기업", "발행일", "애널리스트"]
    values = [
        (info.get("benchmark") or meta_dict.get("benchmark", ""))[:30],
        f"{n_tickers}개 종목",
        today,
        "Lens AI",
    ]
    col_w = Inches(1.85)
    for i, (h, v) in enumerate(zip(headers, values)):
        hc = tbl2.rows[0].cells[i]
        vc = tbl2.rows[1].cells[i]
        hc.width = vc.width = col_w
        _set_cell_bg(hc, DARK)
        ph = hc.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = ph.add_run(h)
        rh.font.name = "Arial"; rh.font.bold = True
        rh.font.size = Pt(8); rh.font.color.rgb = _rgb(WHITE)
        pv = vc.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(v)
        rv.font.name = "Arial"; rv.font.size = Pt(8)
        rv.font.color.rgb = _rgb(GRAY2)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 8) LENS 슬로건 푸터
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(4)
    rf = pf.add_run('"선명하게, 다르게 본다."  Insight in focus.   [ L ]  LENS CAPITAL RESEARCH  │  lens-capital.io')
    rf.font.name = "Arial"; rf.font.italic = True
    rf.font.size = Pt(8); rf.font.color.rgb = _rgb(DARK_NAVY)

    doc.add_paragraph()

    # 9) 목차
    _add_section_heading(doc, "01  목차 (Table of Contents)")
    toc_items = [
        "II.   Executive Summary (투자요약)",
        "III.  현재 시황 & 산업 동향 (Market Context)",
        "IV.   밸류체인 맵 & 수혜/소외 매트릭스 (Value Chain & Beneficiary Map)",
        "V.    핵심 동인 분석 (Key Drivers)",
        "VI.   산업 KPI 대시보드",
        "VII.  Bull/Bear 시나리오 & 종목 워치리스트",
        "VIII. 리스크 요인 (Risk Factors)",
        "IX.   종합 결론 (Conclusion)",
    ]
    for t in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(t)
        r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.color.rgb = _rgb(GRAY2)


# ── 섹션 렌더러 ───────────────────────────────────────────────────────────────

def _add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _add_bottom_border(p, BRAND_RED, "8", "1")
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.bold = True
    r.font.size = Pt(13); r.font.color.rgb = _rgb(DARK_NAVY)


def _add_sub_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.bold = True
    r.font.size = Pt(11); r.font.color.rgb = _rgb(BRAND_RED)


def _add_mixed_run(paragraph, text: str, size: Pt):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.font.bold = True
        else:
            r = paragraph.add_run(part)
        r.font.name = "Arial"; r.font.size = size
        r.font.color.rgb = _rgb(GRAY2)


def _add_table(doc, lines: list):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-| :]+\|\s*$", line):
            continue
        cells = [c.strip().strip("*") for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    col_w = 9360 // col_count  # DXA

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Pt(col_w)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.name = "Arial"
            if i == 0:
                r.font.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = _rgb(WHITE)
                _set_cell_bg(cell, DARK_NAVY)
            else:
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(GRAY2)
                if i % 2 == 0:
                    _set_cell_bg(cell, LIGHT_GRAY)
    doc.add_paragraph()


def _add_sections(doc, raw: str):
    SKIP_SECTIONS = {"HEADER", "header"}
    lines = raw.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading not in SKIP_SECTIONS:
                _add_section_heading(doc, heading)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        if stripped.startswith("### "):
            _add_sub_heading(doc, stripped[4:])
            i += 1
            continue

        if re.match(r"^\*\*[^*]+\*\*$", stripped) and len(stripped) > 4:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped.strip("*"))
            r.font.name = "Arial"; r.font.bold = True
            r.font.size = Pt(11); r.font.color.rgb = _rgb(DARK_NAVY)
            i += 1
            continue

        if stripped.startswith(("① ", "② ", "③ ", "④ ", "⑤ ")):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(stripped)
            r.font.name = "Arial"; r.font.bold = True
            r.font.size = Pt(11); r.font.color.rgb = _rgb(DARK_NAVY)
            i += 1
            continue

        if stripped.startswith(("● HIGH", "● MED", "● LOW")):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(stripped)
            r.font.name = "Arial"; r.font.bold = True; r.font.size = Pt(11)
            level = "HIGH" if "HIGH" in stripped else "MED" if "MED" in stripped else "LOW"
            color = "C62828" if level == "HIGH" else "E65100" if level == "MED" else "2E7D32"
            r.font.color.rgb = _rgb(color)
            i += 1
            continue

        if stripped.startswith(("• ", "* ", "▸", "*▸")):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            clean = re.sub(r"^[•*▸ ]+", "", stripped)
            _add_mixed_run(p, clean, Pt(10))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_mixed_run(p, stripped, Pt(10))
        i += 1


# ── 메인 생성 함수 ────────────────────────────────────────────────────────────

def generate_industry_docx(
    sections: dict,
    name_kr: str,
    name_en: str,
    meta_dict: dict,
    output_dir: str,
) -> str:
    raw = sections.get("_raw", "")
    info = _parse_industry_header(raw)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    _add_industry_cover(doc, name_kr, name_en, meta_dict, info)
    doc.add_page_break()
    _add_sections(doc, raw)

    # 마지막 면책 고지
    doc.add_paragraph()
    p_disc = doc.add_paragraph()
    p_disc.paragraph_format.space_before = Pt(20)
    r_disc = p_disc.add_run(
        "본 리포트는 LENS CAPITAL RESEARCH가 작성한 산업 분석 자료로, "
        "[E] 표기 수치는 추정치입니다. 투자 권유를 목적으로 하지 않습니다."
    )
    r_disc.font.name = "Arial"; r_disc.font.italic = True
    r_disc.font.size = Pt(8); r_disc.font.color.rgb = _rgb(GRAY)

    safe_name = re.sub(r"[^\w]", "_", name_kr)[:30]
    filename = f"LENS_Industry_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"[Industry DOCX] 저장 완료: {filepath}")
    return filepath


def generate_industry_pdf(
    sections: dict,
    name_kr: str,
    name_en: str,
    meta_dict: dict,
    output_dir: str,
) -> str:
    docx_path = generate_industry_docx(sections, name_kr, name_en, meta_dict, output_dir)
    pdf_path = docx_path.replace(".docx", ".pdf")

    system = platform.system()
    lo_cmds = []
    if system == "Windows":
        lo_cmds = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
    elif system == "Darwin":
        lo_cmds = ["/Applications/LibreOffice.app/Contents/MacOS/soffice"]
    else:
        lo_cmds = ["libreoffice", "soffice"]

    for lo in lo_cmds:
        try:
            result = subprocess.run(
                [lo, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
                capture_output=True,
                timeout=60,
            )
            if result.returncode == 0 and os.path.exists(pdf_path):
                print(f"[Industry PDF] LibreOffice 변환 완료: {pdf_path}")
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            print(f"[Industry PDF] docx2pdf 변환 완료: {pdf_path}")
            return pdf_path
    except ImportError:
        pass
    except Exception as e:
        print(f"[Industry PDF] docx2pdf 오류: {e}")

    print("[Industry PDF] 변환 실패 — DOCX로 대체")
    return docx_path
